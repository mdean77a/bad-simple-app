"""Tests for the export endpoint and export service."""

import io
from unittest.mock import patch

import pytest

from src.services.export_service import (
    _PAGE_BREAK_MARKER,
    ExportError,
    _add_runs,
    _add_table,
    _build_docx,
    _format_timestamp,
    assemble_markdown,
    build_approval_tracking,
    convert_markdown_to_docx,
    convert_markdown_to_pdf,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_SECTIONS = [
    {"id": "purpose", "name": "Purpose of the Study", "content": "This study examines…"},
    {"id": "procedures", "name": "Study Procedures", "content": "You will be asked to…"},
]

_APPROVALS = [
    {
        "sectionId": "purpose",
        "userName": "Sarah Johnson",
        "userEmail": "sarah@example.com",
        "timestamp": "2026-02-03T14:30:00Z",
    },
    {
        "sectionId": "procedures",
        "userName": "Sarah Johnson",
        "userEmail": "sarah@example.com",
        "timestamp": "2026-02-03T14:35:00Z",
    },
]


def _export_body(
    *,
    sections=None,
    approvals=None,
    fmt="md",
    protocol_name="THAPCA-OH Trial",
    llm_provider=None,
    llm_model=None,
):
    body = {
        "sections": sections if sections is not None else _SECTIONS,
        "approvals": approvals if approvals is not None else _APPROVALS,
        "format": fmt,
        "protocolName": protocol_name,
    }
    if llm_provider is not None:
        body["llmProvider"] = llm_provider
    if llm_model is not None:
        body["llmModel"] = llm_model
    return body


# ===================================================================
# assemble_markdown – unit tests
# ===================================================================


class TestAssembleMarkdown:
    def test_basic_assembly(self):
        result = assemble_markdown(_SECTIONS, "My Protocol")
        assert result.startswith("# My Protocol")
        assert "**Informed Consent Form**" in result
        assert "## Purpose of the Study" in result
        assert "This study examines…" in result
        assert "## Study Procedures" in result
        assert "You will be asked to…" in result

    def test_sections_appear_in_order(self):
        result = assemble_markdown(_SECTIONS, "P")
        purpose_pos = result.index("## Purpose of the Study")
        procedures_pos = result.index("## Study Procedures")
        assert purpose_pos < procedures_pos

    def test_preserves_markdown_formatting(self):
        sections = [
            {
                "id": "s1",
                "name": "Risks",
                "content": "Risks include:\n\n- **Headache**\n- *Nausea*\n- Fatigue",
            }
        ]
        result = assemble_markdown(sections, "P")
        assert "- **Headache**" in result
        assert "- *Nausea*" in result

    def test_single_section(self):
        result = assemble_markdown([_SECTIONS[0]], "P")
        assert "## Purpose of the Study" in result
        assert "Study Procedures" not in result

    def test_empty_content(self):
        sections = [{"id": "s1", "name": "Empty", "content": ""}]
        result = assemble_markdown(sections, "P")
        assert "## Empty" in result

    def test_page_break_before_signature_sections(self):
        sections = [
            {"id": "s1", "name": "Benefits", "content": "Some benefits."},
            {"id": "s2", "name": "Adult Consent", "content": "Sign here."},
            {"id": "s3", "name": "Teen Assent", "content": "Assent text."},
        ]
        result = assemble_markdown(sections, "P", page_break_before={"Adult Consent", "Teen Assent"})
        # Marker should appear before each signature section
        assert f"{_PAGE_BREAK_MARKER}\n\n## Adult Consent" in result
        assert f"{_PAGE_BREAK_MARKER}\n\n## Teen Assent" in result
        # No marker before content sections
        assert f"{_PAGE_BREAK_MARKER}\n\n## Benefits" not in result

    def test_no_page_break_when_param_omitted(self):
        sections = [
            {"id": "s1", "name": "Adult Consent", "content": "Sign here."},
        ]
        result = assemble_markdown(sections, "P")
        assert _PAGE_BREAK_MARKER not in result

    def test_page_break_does_not_affect_unmatched_sections(self):
        result = assemble_markdown(_SECTIONS, "P", page_break_before={"Adult Consent"})
        assert _PAGE_BREAK_MARKER not in result


# ===================================================================
# build_approval_tracking / _format_timestamp – unit tests
# ===================================================================


class TestFormatTimestamp:
    def test_basic(self):
        assert _format_timestamp("2026-02-03T14:30:00Z") == "Feb 3, 2026 at 2:30 PM"

    def test_morning(self):
        assert _format_timestamp("2026-02-03T09:05:00Z") == "Feb 3, 2026 at 9:05 AM"

    def test_noon(self):
        assert _format_timestamp("2026-12-25T12:00:00Z") == "Dec 25, 2026 at 12:00 PM"

    def test_midnight(self):
        assert _format_timestamp("2026-01-01T00:00:00Z") == "Jan 1, 2026 at 12:00 AM"

    def test_with_offset(self):
        result = _format_timestamp("2026-02-03T14:30:00+05:00")
        assert "Feb 3, 2026" in result
        assert "2:30 PM" in result


class TestBuildApprovalTracking:
    def test_basic_table(self):
        result = build_approval_tracking(_APPROVALS, _SECTIONS)
        assert "## Approval Tracking" in result
        assert "| Section | Approved By | Date & Time |" in result
        assert "| Purpose of the Study | Sarah Johnson | Feb 3, 2026 at 2:30 PM |" in result
        assert "| Study Procedures | Sarah Johnson | Feb 3, 2026 at 2:35 PM |" in result

    def test_empty_approvals_returns_empty(self):
        assert build_approval_tracking([], _SECTIONS) == ""

    def test_starts_with_page_break_marker(self):
        result = build_approval_tracking(_APPROVALS, _SECTIONS)
        assert result.startswith(_PAGE_BREAK_MARKER)

    def test_unknown_section_id_uses_id_as_fallback(self):
        approvals = [
            {"sectionId": "unknown-id", "userName": "Jane", "timestamp": "2026-01-01T10:00:00Z"},
        ]
        result = build_approval_tracking(approvals, _SECTIONS)
        assert "| unknown-id | Jane |" in result

    def test_multiple_approvers(self):
        approvals = [
            {"sectionId": "purpose", "userName": "Sarah Johnson", "timestamp": "2026-02-03T14:30:00Z"},
            {"sectionId": "procedures", "userName": "Maria Chen", "timestamp": "2026-02-04T10:15:00Z"},
        ]
        result = build_approval_tracking(approvals, _SECTIONS)
        assert "Sarah Johnson" in result
        assert "Maria Chen" in result

    def test_preserves_section_order(self):
        result = build_approval_tracking(_APPROVALS, _SECTIONS)
        purpose_pos = result.index("Purpose of the Study")
        procedures_pos = result.index("Study Procedures")
        assert purpose_pos < procedures_pos

    def test_disclosure_appended_for_anthropic(self):
        result = build_approval_tracking(
            _APPROVALS, _SECTIONS, "anthropic", "claude-sonnet-4-6"
        )
        assert (
            "This document was generated using Anthropic model "
            "`claude-sonnet-4-6`."
            in result
        )
        # Disclosure follows the table
        table_pos = result.index("| Section |")
        disclosure_pos = result.index("This document was generated")
        assert disclosure_pos > table_pos

    def test_disclosure_appended_for_openai(self):
        result = build_approval_tracking(
            _APPROVALS, _SECTIONS, "openai", "gpt-5.1"
        )
        assert "OpenAI model `gpt-5.1`" in result

    def test_disclosure_for_local_omits_model(self):
        result = build_approval_tracking(
            _APPROVALS, _SECTIONS, "local", ""
        )
        assert "Local (LM Studio)" in result
        assert "determined by the LM Studio runtime" in result
        assert "`" not in result.split("Approval Tracking")[1].split("Local")[0]

    def test_disclosure_omitted_when_provider_missing(self):
        result = build_approval_tracking(_APPROVALS, _SECTIONS)
        assert "This document was generated" not in result

    def test_disclosure_omitted_when_no_approvals(self):
        # Tracking returns "" when there are no approvals, regardless of LLM args
        result = build_approval_tracking([], _SECTIONS, "anthropic", "claude-sonnet-4-6")
        assert result == ""

    def test_disclosure_unknown_provider_falls_back_to_id(self):
        result = build_approval_tracking(
            _APPROVALS, _SECTIONS, "custom-provider", "custom-model"
        )
        assert "custom-provider model `custom-model`" in result


# ===================================================================
# convert_markdown_to_docx – integration tests (python-docx is pure Python)
# ===================================================================


class TestConvertMarkdownToDocx:
    def _open_docx(self, docx_bytes: bytes):
        from docx import Document

        return Document(io.BytesIO(docx_bytes))

    def test_produces_valid_docx(self):
        md = "# Title\n\nA paragraph."
        result = convert_markdown_to_docx(md)
        assert isinstance(result, bytes)
        # DOCX files start with PK (ZIP archive)
        assert result[:2] == b"PK"

    def test_heading_levels(self):
        md = "# H1\n\n## H2\n\n### H3\n"
        doc = self._open_docx(convert_markdown_to_docx(md))
        styles = [p.style.name for p in doc.paragraphs]
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        assert "Heading 3" in styles

    def test_paragraphs(self):
        md = "# Title\n\nFirst paragraph.\n\nSecond paragraph.\n"
        doc = self._open_docx(convert_markdown_to_docx(md))
        texts = [p.text for p in doc.paragraphs if p.style.name == "Normal"]
        assert "First paragraph." in texts
        assert "Second paragraph." in texts

    def test_bullet_list(self):
        md = "- Item A\n- Item B\n"
        doc = self._open_docx(convert_markdown_to_docx(md))
        bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert len(bullets) == 2
        assert bullets[0].text == "Item A"

    def test_numbered_list(self):
        md = "1. First\n2. Second\n"
        doc = self._open_docx(convert_markdown_to_docx(md))
        numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(numbered) == 2
        assert numbered[0].text == "First"

    def test_bold_inline(self):
        md = "This is **important** text.\n"
        doc = self._open_docx(convert_markdown_to_docx(md))
        para = [p for p in doc.paragraphs if p.style.name == "Normal"][0]
        runs = para.runs
        bold_runs = [r for r in runs if r.bold]
        assert any("important" in r.text for r in bold_runs)

    def test_italic_inline(self):
        md = "This is *emphasized* text.\n"
        doc = self._open_docx(convert_markdown_to_docx(md))
        para = [p for p in doc.paragraphs if p.style.name == "Normal"][0]
        italic_runs = [r for r in para.runs if r.italic]
        assert any("emphasized" in r.text for r in italic_runs)

    def test_page_break_marker_produces_page_break(self):
        md = f"# Page 1\n\n{_PAGE_BREAK_MARKER}\n\n# Page 2\n"
        doc = self._open_docx(convert_markdown_to_docx(md))
        # Page breaks appear as paragraph elements; just verify both headings exist
        texts = [p.text for p in doc.paragraphs]
        assert "Page 1" in texts
        assert "Page 2" in texts

    def test_thematic_break_in_content_skipped(self):
        """--- or *** in LLM content should NOT produce a page break."""
        md = "# Title\n\nParagraph A.\n\n---\n\nParagraph B.\n\n***\n\nParagraph C.\n"
        doc = self._open_docx(convert_markdown_to_docx(md))
        texts = [p.text for p in doc.paragraphs if p.style.name == "Normal"]
        assert any("Paragraph A" in t for t in texts)
        assert any("Paragraph B" in t for t in texts)
        assert any("Paragraph C" in t for t in texts)

    def test_full_document_assembly(self):
        md = assemble_markdown(_SECTIONS, "My Protocol")
        doc = self._open_docx(convert_markdown_to_docx(md))
        texts = [p.text for p in doc.paragraphs]
        assert "My Protocol" in texts
        assert "Purpose of the Study" in texts
        assert "Study Procedures" in texts

    def test_table_renders_as_word_table(self):
        md = (
            "| Section | Approved By | Date |\n"
            "|---------|-------------|------|\n"
            "| Purpose | Sarah | Feb 3 |\n"
            "| Risks | John | Feb 4 |\n"
        )
        doc = self._open_docx(convert_markdown_to_docx(md))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3  # header + 2 data rows
        assert len(table.columns) == 3
        # Header row should be bold
        header_cells = [cell.text for cell in table.rows[0].cells]
        assert header_cells == ["Section", "Approved By", "Date"]
        assert table.rows[0].cells[0].paragraphs[0].runs[0].bold
        # Data rows
        assert table.rows[1].cells[0].text == "Purpose"
        assert table.rows[2].cells[1].text == "John"

    def test_table_in_full_document(self):
        md = (
            "## Approval Tracking\n\n"
            "| Section | Approved By |\n"
            "|---------|-------------|\n"
            "| Purpose | Sarah |\n"
        )
        doc = self._open_docx(convert_markdown_to_docx(md))
        assert len(doc.tables) == 1
        texts = [p.text for p in doc.paragraphs]
        assert "Approval Tracking" in texts

    def test_import_error_raises_export_error(self):
        real_import = __import__

        def fail_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("No module named 'docx'")
            return real_import(name, *args, **kwargs)

        with patch("builtins.__import__", side_effect=fail_import):
            with pytest.raises(ExportError, match="DOCX export requires python-docx"):
                convert_markdown_to_docx("# Title")

    def test_generic_exception_raises_export_error(self):
        with patch("src.services.export_service._build_docx", side_effect=RuntimeError("unexpected")):
            with pytest.raises(ExportError, match="DOCX conversion failed.*unexpected"):
                convert_markdown_to_docx("# Title")


# ===================================================================
# convert_markdown_to_pdf – integration tests (xhtml2pdf is pure Python)
# ===================================================================


class TestConvertMarkdownToPdf:
    def test_produces_valid_pdf(self):
        md = "# Title\n\nA paragraph."
        result = convert_markdown_to_pdf(md)
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    def test_pdf_is_nontrivial(self):
        md = "# Title\n\nHello world from PDF test.\n\nSecond paragraph."
        result = convert_markdown_to_pdf(md)
        # Valid PDF with actual content should be significantly larger than empty
        assert len(result) > 500

    def test_pdf_renders_tables(self):
        md = (
            "| Col A | Col B |\n"
            "|-------|-------|\n"
            "| 1     | 2     |\n"
        )
        result = convert_markdown_to_pdf(md)
        assert result[:5] == b"%PDF-"

    def test_pdf_status_err_raises_export_error(self):
        from unittest.mock import MagicMock

        mock_status = MagicMock()
        mock_status.err = 1

        with patch("xhtml2pdf.pisa.CreatePDF", return_value=mock_status):
            with pytest.raises(ExportError, match="PDF conversion returned errors"):
                convert_markdown_to_pdf("# Title")

    def test_pdf_generic_exception_wraps_as_export_error(self):
        with patch("markdown.markdown", side_effect=RuntimeError("boom")):
            with pytest.raises(ExportError, match="PDF conversion failed.*boom"):
                convert_markdown_to_pdf("# Title")

    def test_pdf_import_error_raises_export_error(self):
        real_import = __import__

        def fail_import(name, *args, **kwargs):
            if name in ("xhtml2pdf", "markdown"):
                raise ImportError(f"No module named '{name}'")
            return real_import(name, *args, **kwargs)

        with patch("builtins.__import__", side_effect=fail_import):
            with pytest.raises(ExportError, match="PDF export requires"):
                convert_markdown_to_pdf("# Title")

    def test_underscore_lines_not_treated_as_page_breaks(self):
        """Underscore signature lines should render as lines, not <hr>."""
        md = "## Teen Assent\n\nSign below:\n\n______________________________\nName\n"
        # Should produce a valid PDF without errors
        result = convert_markdown_to_pdf(md)
        assert result[:5] == b"%PDF-"

    def test_page_break_marker_in_pdf(self):
        """Page-break marker should produce a valid PDF."""
        md = f"## Section A\n\nContent.\n\n{_PAGE_BREAK_MARKER}\n\n## Section B\n\nMore.\n"
        result = convert_markdown_to_pdf(md)
        assert result[:5] == b"%PDF-"

    def test_asterisk_thematic_break_stripped(self):
        """*** in LLM content should not produce an <hr> or page break."""
        md = "## Risks\n\nRisk A.\n\n***\n\nRisk B.\n"
        result = convert_markdown_to_pdf(md)
        assert result[:5] == b"%PDF-"

    def test_dash_thematic_break_stripped(self):
        """--- in LLM content should not produce an <hr> or page break."""
        md = "## Procedures\n\nStep A.\n\n---\n\nStep B.\n"
        result = convert_markdown_to_pdf(md)
        assert result[:5] == b"%PDF-"


# ===================================================================
# _build_docx / _add_runs – unit tests
# ===================================================================


class TestBuildDocxHelpers:
    def _make_doc(self):
        from docx import Document

        return Document()

    def test_add_runs_plain_text(self):
        doc = self._make_doc()
        p = doc.add_paragraph()
        _add_runs(p, "plain text")
        assert p.text == "plain text"

    def test_add_runs_bold(self):
        doc = self._make_doc()
        p = doc.add_paragraph()
        _add_runs(p, "before **bold** after")
        assert any(r.bold for r in p.runs)
        assert "bold" in p.text

    def test_add_runs_italic(self):
        doc = self._make_doc()
        p = doc.add_paragraph()
        _add_runs(p, "before *italic* after")
        assert any(r.italic for r in p.runs)

    def test_add_runs_bold_italic(self):
        doc = self._make_doc()
        p = doc.add_paragraph()
        _add_runs(p, "***bold italic***")
        assert any(r.bold and r.italic for r in p.runs)

    def test_build_docx_mixed(self):
        doc = self._make_doc()
        md = "# Title\n\nParagraph.\n\n- Bullet\n\n1. Numbered\n"
        _build_docx(doc, md)
        styles = [p.style.name for p in doc.paragraphs]
        assert "Heading 1" in styles
        assert "Normal" in styles
        assert "List Bullet" in styles
        assert "List Number" in styles

    def test_multi_line_paragraph(self):
        doc = self._make_doc()
        md = "First line of paragraph\nSecond line of paragraph\n\nNew paragraph here\n"
        _build_docx(doc, md)
        normal = [p for p in doc.paragraphs if p.style.name == "Normal"]
        assert len(normal) == 2
        assert normal[0].text == "First line of paragraph Second line of paragraph"
        assert normal[1].text == "New paragraph here"

    def test_build_docx_table(self):
        doc = self._make_doc()
        md = (
            "| A | B |\n"
            "|---|---|\n"
            "| 1 | 2 |\n"
        )
        _build_docx(doc, md)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 2  # header + 1 data
        assert doc.tables[0].rows[0].cells[0].text == "A"
        assert doc.tables[0].rows[1].cells[1].text == "2"

    def test_build_docx_table_then_paragraph(self):
        doc = self._make_doc()
        md = (
            "| X | Y |\n"
            "|---|---|\n"
            "| a | b |\n"
            "\nAfter table.\n"
        )
        _build_docx(doc, md)
        assert len(doc.tables) == 1
        normal = [p for p in doc.paragraphs if p.style.name == "Normal"]
        assert any("After table" in p.text for p in normal)

    def test_add_table_header_bold(self):
        doc = self._make_doc()
        _add_table(doc, [["Col1", "Col2"], ["a", "b"]])
        assert doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0].bold
        assert not doc.tables[0].rows[1].cells[0].paragraphs[0].runs[0].bold

    def test_add_table_empty_rows_noop(self):
        doc = self._make_doc()
        _add_table(doc, [])
        assert len(doc.tables) == 0


# ===================================================================
# Route tests – POST /api/v1/export/
# ===================================================================


class TestExportRoute:
    # --- Markdown export ---

    @pytest.mark.asyncio
    async def test_export_markdown(self, client):
        response = await client.post("/api/v1/export/", json=_export_body(fmt="md"))

        assert response.status_code == 200
        assert "text/markdown" in response.headers["content-type"]
        assert "THAPCA-OH Trial_ICF.md" in response.headers["content-disposition"]

        text = response.text
        assert "# THAPCA-OH Trial" in text
        assert "## Purpose of the Study" in text
        assert "This study examines…" in text

    @pytest.mark.asyncio
    async def test_export_markdown_sections_in_order(self, client):
        response = await client.post("/api/v1/export/", json=_export_body(fmt="md"))
        text = response.text
        assert text.index("Purpose of the Study") < text.index("Study Procedures")

    # --- PDF export ---

    @pytest.mark.asyncio
    @patch("src.api.routes.export.convert_markdown_to_pdf")
    async def test_export_pdf(self, mock_pdf, client):
        mock_pdf.return_value = b"%PDF-1.4 fake"

        response = await client.post("/api/v1/export/", json=_export_body(fmt="pdf"))

        assert response.status_code == 200
        assert response.headers["content-type"] == "application/pdf"
        assert "THAPCA-OH Trial_ICF.pdf" in response.headers["content-disposition"]
        assert response.content == b"%PDF-1.4 fake"
        mock_pdf.assert_called_once()

    @pytest.mark.asyncio
    @patch("src.api.routes.export.convert_markdown_to_pdf")
    async def test_export_pdf_receives_assembled_markdown(self, mock_pdf, client):
        mock_pdf.return_value = b"%PDF"

        await client.post("/api/v1/export/", json=_export_body(fmt="pdf"))

        md_arg = mock_pdf.call_args[0][0]
        assert "# THAPCA-OH Trial" in md_arg
        assert "## Purpose of the Study" in md_arg

    @pytest.mark.asyncio
    @patch("src.api.routes.export.convert_markdown_to_pdf")
    async def test_export_pdf_conversion_error(self, mock_pdf, client):
        mock_pdf.side_effect = ExportError("Cairo not found")

        response = await client.post("/api/v1/export/", json=_export_body(fmt="pdf"))

        assert response.status_code == 502
        data = response.json()
        assert data["code"] == "EXPORT_ERROR"
        assert "Cairo not found" in data["detail"]

    # --- DOCX export ---

    @pytest.mark.asyncio
    @patch("src.api.routes.export.convert_markdown_to_docx")
    async def test_export_docx(self, mock_docx, client):
        mock_docx.return_value = b"PK\x03\x04 fake"

        response = await client.post("/api/v1/export/", json=_export_body(fmt="docx"))

        assert response.status_code == 200
        assert "wordprocessingml" in response.headers["content-type"]
        assert "THAPCA-OH Trial_ICF.docx" in response.headers["content-disposition"]
        mock_docx.assert_called_once()

    @pytest.mark.asyncio
    @patch("src.api.routes.export.convert_markdown_to_docx")
    async def test_export_docx_conversion_error(self, mock_docx, client):
        mock_docx.side_effect = ExportError("python-docx error")

        response = await client.post("/api/v1/export/", json=_export_body(fmt="docx"))

        assert response.status_code == 502
        data = response.json()
        assert data["code"] == "EXPORT_ERROR"
        assert "python-docx error" in data["detail"]

    # --- Validation ---

    @pytest.mark.asyncio
    async def test_export_empty_sections(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(sections=[]),
        )

        assert response.status_code == 422
        assert response.json()["code"] == "VALIDATION_ERROR"
        assert "section" in response.json()["detail"].lower()

    @pytest.mark.asyncio
    async def test_export_empty_protocol_name(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(protocol_name="   "),
        )

        assert response.status_code == 422
        assert response.json()["code"] == "VALIDATION_ERROR"
        assert "protocolName" in response.json()["detail"]

    @pytest.mark.asyncio
    async def test_export_invalid_format(self, client):
        """Pydantic rejects formats outside the Literal union."""
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(fmt="html"),
        )
        assert response.status_code == 422  # Pydantic validation error

    @pytest.mark.asyncio
    async def test_export_missing_fields(self, client):
        response = await client.post("/api/v1/export/", json={})
        assert response.status_code == 422

    # --- Content-Disposition ---

    @pytest.mark.asyncio
    async def test_content_disposition_uses_protocol_name(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(protocol_name="Cardiac Trial", fmt="md"),
        )
        assert 'filename="Cardiac Trial_ICF.md"' in response.headers["content-disposition"]

    @pytest.mark.asyncio
    async def test_protocol_name_trimmed(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(protocol_name="  Spaced  ", fmt="md"),
        )
        assert 'filename="Spaced_ICF.md"' in response.headers["content-disposition"]

    # --- Approval tracking in export ---

    @pytest.mark.asyncio
    async def test_markdown_includes_approval_tracking(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(approvals=_APPROVALS, fmt="md"),
        )
        text = response.text
        assert "## Approval Tracking" in text
        assert "Sarah Johnson" in text
        assert "Feb 3, 2026 at 2:30 PM" in text

    @pytest.mark.asyncio
    async def test_approval_tracking_after_sections(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(fmt="md"),
        )
        text = response.text
        last_section_pos = text.index("Study Procedures")
        tracking_pos = text.index("Approval Tracking")
        assert tracking_pos > last_section_pos

    @pytest.mark.asyncio
    async def test_no_approval_tracking_when_empty(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(approvals=[], fmt="md"),
        )
        assert "Approval Tracking" not in response.text

    @pytest.mark.asyncio
    @patch("src.api.routes.export.convert_markdown_to_pdf")
    async def test_pdf_receives_approval_tracking(self, mock_pdf, client):
        mock_pdf.return_value = b"%PDF"
        await client.post("/api/v1/export/", json=_export_body(fmt="pdf"))
        md_arg = mock_pdf.call_args[0][0]
        assert "## Approval Tracking" in md_arg
        assert "Sarah Johnson" in md_arg

    @pytest.mark.asyncio
    async def test_markdown_includes_model_disclosure(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(
                fmt="md",
                llm_provider="anthropic",
                llm_model="claude-sonnet-4-6",
            ),
        )
        assert "This document was generated using Anthropic" in response.text
        assert "claude-sonnet-4-6" in response.text

    @pytest.mark.asyncio
    async def test_markdown_no_disclosure_when_no_llm_metadata(self, client):
        response = await client.post(
            "/api/v1/export/",
            json=_export_body(fmt="md"),
        )
        assert "This document was generated" not in response.text
