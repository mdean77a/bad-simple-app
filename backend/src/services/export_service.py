"""Export service – assemble and convert ICF documents."""

import io
import re
from datetime import datetime


class ExportError(Exception):
    """Raised when export processing fails."""


# Unique marker that cannot appear in LLM-generated content.
# Converters replace it with the format-appropriate page break.
_PAGE_BREAK_MARKER = "<!-- pagebreak -->"


# ---------------------------------------------------------------------------
# Markdown assembly
# ---------------------------------------------------------------------------


def assemble_markdown(
    sections: list[dict[str, str]],
    protocol_name: str,
    page_break_before: set[str] | None = None,
) -> str:
    """Assemble sections into a complete Markdown document.

    Each section dict must have ``name`` and ``content`` keys.
    *page_break_before* is an optional set of section names that should
    be preceded by a page-break marker.
    """
    parts: list[str] = [f"# {protocol_name}", "", "**Informed Consent Form**", ""]

    for section in sections:
        if page_break_before and section["name"] in page_break_before:
            parts.append(_PAGE_BREAK_MARKER)
            parts.append("")
        parts.append(f"## {section['name']}")
        parts.append("")
        parts.append(section["content"])
        parts.append("")

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Approval tracking page
# ---------------------------------------------------------------------------


_PROVIDER_LABELS: dict[str, str] = {
    "anthropic": "Anthropic",
    "openai": "OpenAI",
    "local": "Local (LM Studio)",
}


def _build_model_disclosure(llm_provider: str | None, llm_model: str | None) -> str:
    if not llm_provider:
        return ""
    label = _PROVIDER_LABELS.get(llm_provider, llm_provider)
    if llm_provider == "local":
        return (
            f"This document was generated using {label}; "
            "the specific model is determined by the LM Studio runtime."
        )
    if not llm_model:
        return f"This document was generated using {label}."
    return f"This document was generated using {label} model `{llm_model}`."


def build_approval_tracking(
    approvals: list[dict[str, str]],
    sections: list[dict[str, str]],
    llm_provider: str | None = None,
    llm_model: str | None = None,
) -> str:
    """Build an approval tracking Markdown section.

    Returns an empty string when *approvals* is empty.
    Each approval dict must have ``sectionId``, ``userName``, ``timestamp``.
    Each section dict must have ``id`` and ``name``.
    When *llm_provider* is provided, a disclosure statement is appended
    after the table indicating which provider/model generated the document.
    """
    if not approvals:
        return ""

    section_names: dict[str, str] = {s["id"]: s["name"] for s in sections}

    rows: list[str] = []
    for a in approvals:
        name = section_names.get(a["sectionId"], a["sectionId"])
        ts = _format_timestamp(a["timestamp"])
        rows.append(f"| {name} | {a['userName']} | {ts} |")

    lines = [
        _PAGE_BREAK_MARKER,
        "",
        "## Approval Tracking",
        "",
        "| Section | Approved By | Date & Time |",
        "|---------|-------------|-------------|",
        *rows,
        "",
    ]
    disclosure = _build_model_disclosure(llm_provider, llm_model)
    if disclosure:
        lines.append(disclosure)
        lines.append("")
    return "\n".join(lines)


def _format_timestamp(iso_timestamp: str) -> str:
    """Format an ISO-8601 timestamp as ``Feb 3, 2026 at 2:30 PM``."""
    dt = datetime.fromisoformat(iso_timestamp)
    month = dt.strftime("%b")
    day = dt.day
    year = dt.year
    time_str = dt.strftime("%I:%M %p").lstrip("0")
    return f"{month} {day}, {year} at {time_str}"


# ---------------------------------------------------------------------------
# PDF conversion (xhtml2pdf – pure Python, no system dependencies)
# ---------------------------------------------------------------------------

_PDF_CSS = """\
body {
    font-family: Helvetica, serif;
    font-size: 12pt;
    line-height: 1.6;
}
h1 {
    text-align: center;
    font-size: 20pt;
    margin-bottom: 0.5em;
}
h2 {
    font-size: 16pt;
    margin-top: 1.5em;
    page-break-after: avoid;
}
h3 {
    font-size: 14pt;
    margin-top: 1em;
    page-break-after: avoid;
}
table {
    border-collapse: collapse;
    width: 100%;
    margin: 1em 0;
}
th, td {
    border: 1px solid #999;
    padding: 6px 10px;
    text-align: left;
}
th {
    background-color: #f0f0f0;
}
hr {
    border: none;
    border-top: 1px solid #ccc;
    margin: 1.5em 0;
}
"""


def convert_markdown_to_pdf(md_content: str) -> bytes:
    """Convert Markdown content to PDF bytes."""
    try:
        import markdown as md_lib
        from xhtml2pdf import pisa
    except ImportError as exc:
        raise ExportError(
            f"PDF export requires xhtml2pdf and markdown packages: {exc}"
        ) from exc

    try:
        # 1) Replace our unique page-break marker with an HTML page break.
        md_content = md_content.replace(
            _PAGE_BREAK_MARKER,
            '<div style="page-break-after: always;"></div>',
        )
        # 2) Neutralise ALL Markdown thematic breaks (---, ***, ___) that
        #    may appear in LLM-generated content so the markdown library
        #    does not convert them to <hr> (which would waste vertical
        #    space or trigger unwanted page breaks in some renderers).
        #    Underscore lines are signature lines — render as a visible line.
        md_content = re.sub(
            r"^_{3,}\s*$",
            '<div style="border-bottom: 1px solid black; width: 50%;'
            ' margin: 0.5em 0;"></div>',
            md_content,
            flags=re.MULTILINE,
        )
        #    Dash and asterisk thematic breaks — remove entirely.
        md_content = re.sub(
            r"^[-*]{3,}\s*$",
            "",
            md_content,
            flags=re.MULTILINE,
        )

        html_body = md_lib.markdown(
            md_content,
            extensions=["tables", "fenced_code", "sane_lists"],
        )
        html_doc = (
            f"<html><head>"
            f"<meta charset='utf-8'/>"
            f"<style>{_PDF_CSS}</style>"
            f"</head><body>{html_body}</body></html>"
        )
        buf = io.BytesIO()
        status = pisa.CreatePDF(html_doc, dest=buf)
        if status.err:
            raise ExportError("PDF conversion returned errors")
        return buf.getvalue()
    except ExportError:
        raise
    except Exception as exc:
        raise ExportError(f"PDF conversion failed: {exc}") from exc


# ---------------------------------------------------------------------------
# DOCX conversion (python-docx – lazy-imported for testability)
# ---------------------------------------------------------------------------


def convert_markdown_to_docx(md_content: str) -> bytes:
    """Convert Markdown content to DOCX bytes."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise ExportError(
            f"DOCX export requires python-docx package: {exc}"
        ) from exc

    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        _build_docx(doc, md_content)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ExportError:
        raise
    except Exception as exc:
        raise ExportError(f"DOCX conversion failed: {exc}") from exc


# ---------------------------------------------------------------------------
# Internal helpers (no library imports needed – work on passed objects)
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.+)$")
_HRULE_RE = re.compile(r"^[-*]{3,}\s*$")
_TABLE_ROW_RE = re.compile(r"^\|.+\|$")
_TABLE_SEP_RE = re.compile(r"^\|[-:\s|]+\|$")
_INLINE_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")


def _build_docx(doc, md_content: str) -> None:  # noqa: ANN001
    """Walk Markdown lines and populate a python-docx *Document*."""
    lines = md_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Intentional page-break marker
        if stripped == _PAGE_BREAK_MARKER:
            doc.add_page_break()
            i += 1
            continue

        # Thematic breaks (---, ***) from LLM content — skip silently
        if _HRULE_RE.match(stripped):
            i += 1
            continue

        # Heading
        m = _HEADING_RE.match(stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # Bullet list item
        m = _BULLET_RE.match(stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(1))
            i += 1
            continue

        # Numbered list item
        m = _NUMBERED_RE.match(stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(1))
            i += 1
            continue

        # Markdown table
        if _TABLE_ROW_RE.match(stripped):
            table_rows: list[list[str]] = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i].strip()):
                row_line = lines[i].strip()
                if _TABLE_SEP_RE.match(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            if table_rows:
                _add_table(doc, table_rows)
            continue

        # Regular paragraph – accumulate consecutive non-special lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt == _PAGE_BREAK_MARKER
                or _HEADING_RE.match(nxt)
                or _BULLET_RE.match(nxt)
                or _NUMBERED_RE.match(nxt)
                or _HRULE_RE.match(nxt)
                or _TABLE_ROW_RE.match(nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1

        p = doc.add_paragraph()
        _add_runs(p, " ".join(para_lines))


def _add_table(doc, rows: list[list[str]]) -> None:  # noqa: ANN001
    """Add a Word table from parsed Markdown table rows.

    The first row is treated as the header.
    """
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")
    for r_idx, cells in enumerate(rows):
        for c_idx, cell_text in enumerate(cells):
            if c_idx < n_cols:
                cell = table.rows[r_idx].cells[c_idx]
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(cell_text)
                if r_idx == 0:
                    run.bold = True


def _add_runs(paragraph, text: str) -> None:  # noqa: ANN001
    """Add text with bold / italic inline formatting to a paragraph."""
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
